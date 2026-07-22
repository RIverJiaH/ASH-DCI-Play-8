from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
ASSETS = Path(__file__).resolve().parent / "assets"
OUTPUT = ROOT / "docs" / "脑护通Demo V3用户使用手册.docx"

FONT_LATIN = "Calibri"
FONT_CJK = "Microsoft YaHei"
INK = "17242B"
MUTED = "63727A"
TEAL = "176B68"
TEAL_DARK = "0D4F4D"
TEAL_SOFT = "E4F1F0"
BLUE = "2E74B5"
BLUE_DARK = "1F4D78"
BLUE_GRAY = "E8EEF5"
AMBER = "9A5A08"
AMBER_SOFT = "FFF3D9"
RED = "B4232F"
RED_SOFT = "FBEAEC"
LINE = "D7E0E3"
WHITE = "FFFFFF"


def set_run(run, size=None, bold=None, color=None, italic=None):
    run.font.name = FONT_LATIN
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    fonts.set(qn("w:ascii"), FONT_LATIN)
    fonts.set(qn("w:hAnsi"), FONT_LATIN)
    fonts.set(qn("w:eastAsia"), FONT_CJK)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def set_paragraph_fill(paragraph, fill):
    ppr = paragraph._p.get_or_add_pPr()
    shading = ppr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        ppr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_paragraph_border(paragraph, edge="left", color=TEAL, size="18", space="8"):
    ppr = paragraph._p.get_or_add_pPr()
    borders = ppr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        ppr.append(borders)
    node = borders.find(qn(f"w:{edge}"))
    if node is None:
        node = OxmlElement(f"w:{edge}")
        borders.append(node)
    node.set(qn("w:val"), "single")
    node.set(qn("w:sz"), size)
    node.set(qn("w:space"), space)
    node.set(qn("w:color"), color)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_end])
    set_run(run, 8, color=MUTED)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.start_type = WD_SECTION_START.NEW_PAGE

    normal = doc.styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, BLUE_DARK, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = doc.styles["Caption"]
    caption.font.name = FONT_LATIN
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    caption.font.size = Pt(9)
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    configure_header_footer(section)
    return configure_numbering(doc)


def configure_header_footer(section):
    header = section.header
    hp = header.paragraphs[0]
    hp.paragraph_format.space_after = Pt(0)
    hp.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
    set_run(hp.add_run("脑护通  |  AI临床情境辅助 Demo V3"), 8, True, TEAL_DARK)
    set_run(hp.add_run("\t用户使用手册"), 8, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.paragraph_format.space_after = Pt(0)
    fp.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
    set_run(fp.add_run("演示版本  |  2026.07"), 8, color=MUTED)
    set_run(fp.add_run("\t第 "), 8, color=MUTED)
    add_page_field(fp)
    set_run(fp.add_run(" 页"), 8, color=MUTED)


def configure_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    next_abstract = max(abstract_ids, default=0) + 1
    next_num = max(num_ids, default=0) + 1

    def add_definition(abstract_id, num_id, fmt, text, font=None):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        level = OxmlElement("w:lvl")
        level.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        ppr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        indent = OxmlElement("w:ind")
        indent.set(qn("w:left"), "540")
        indent.set(qn("w:hanging"), "270")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        ppr.extend([tabs, indent, spacing])
        level.extend([start, num_fmt, lvl_text, suff, ppr])
        if font:
            rpr = OxmlElement("w:rPr")
            fonts = OxmlElement("w:rFonts")
            fonts.set(qn("w:ascii"), font)
            fonts.set(qn("w:hAnsi"), font)
            rpr.append(fonts)
            level.append(rpr)
        abstract.append(level)
        numbering.append(abstract)

        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(abstract_id))
        num.append(abstract_ref)
        numbering.append(num)

    add_definition(next_abstract, next_num, "decimal", "%1.")
    add_definition(next_abstract + 1, next_num + 1, "bullet", "•", FONT_LATIN)
    return {"number": next_num, "bullet": next_num + 1}


def add_list_item(doc, text, num_id, bold_prefix=None):
    p = doc.add_paragraph()
    ppr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    ppr.insert(0, num_pr)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        set_run(p.add_run(bold_prefix), 11, True, INK)
        set_run(p.add_run(text[len(bold_prefix):]), 11, color=INK)
    else:
        set_run(p.add_run(text), 11, color=INK)
    return p


def add_kicker(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    set_run(p.add_run(text.upper()), 9, True, TEAL_DARK)
    return p


def add_title(doc, title, subtitle=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    set_run(p.add_run(title), 28, True, INK)
    if subtitle:
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(14)
        set_run(sp.add_run(subtitle), 14, color=TEAL_DARK)


def add_callout(doc, label, text, fill=TEAL_SOFT, accent=TEAL):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(9)
    p.paragraph_format.line_spacing = 1.25
    set_paragraph_fill(p, fill)
    set_paragraph_border(p, "left", accent)
    set_run(p.add_run(f"{label}  "), 10.5, True, accent)
    set_run(p.add_run(text), 10.5, color=INK)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        set_run(p.add_run(bold_prefix), 11, True, INK)
        set_run(p.add_run(text[len(bold_prefix):]), 11, color=INK)
    else:
        set_run(p.add_run(text), 11, color=INK)
    return p


def add_screenshot(doc, filename, caption, alt_text):
    path = ASSETS / filename
    if not path.exists():
        raise FileNotFoundError(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    shape = run.add_picture(str(path), width=Inches(6.5))
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", alt_text)
    set_paragraph_border(p, "top", LINE, "6", "3")
    set_paragraph_border(p, "bottom", LINE, "6", "3")
    cap = doc.add_paragraph(style="Caption")
    cap.add_run(caption)


def page_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_break(WD_BREAK.PAGE)


def build_manual():
    doc = Document()
    numbering = configure_document(doc)

    # Page 1: editorial cover.
    for _ in range(3):
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(20)
    add_kicker(doc, "USER GUIDE  |  VERSION 3")
    add_title(doc, "脑护通 Demo V3\n用户使用手册", "AI临床情境辅助 · 受控脑控表达与护理任务闭环")
    add_callout(
        doc,
        "适用范围",
        "本手册用于第三版功能演示。患者输入、病历和脑控信号均为模拟数据，不用于诊断、治疗或替代临床护理流程。",
        AMBER_SOFT,
        AMBER,
    )
    add_body(doc, "文档对象：客户演示人员、项目成员和护理端体验人员")
    add_body(doc, "版本日期：2026年7月")
    page_break(doc)

    # Page 2: quick start.
    doc.add_heading("1  快速启动", level=1)
    add_callout(doc, "启动文件", r"C:\Users\19340\Documents\Codex\2026-07-21\bang\brain-care-demo\start_public_demo.bat")
    for text in (
        "双击 start_public_demo.bat，等待黑色命令窗口显示本地网址。",
        "浏览器打开 http://127.0.0.1:8000/。如 cpolar 正常，黑色窗口还会显示临时公网网址。",
        "演示期间保持黑色命令窗口开启；关闭窗口后网站停止运行。",
        "升级代码后必须关闭旧窗口并重新运行 BAT，确保服务端加载最新版本。",
    ):
        add_list_item(doc, text, numbering["number"])
    doc.add_heading("启动成功标准", level=2)
    for text in (
        "页面左上角显示“AI临床情境辅助 Demo V3”。",
        "患者端与护理端可以切换。",
        "右侧显示“模拟脑控信号”和“模拟病历场景”。",
    ):
        add_list_item(doc, text, numbering["bullet"])
    add_callout(doc, "重要", "公网穿透不可用不影响本机演示。先使用本地网址确认页面和后端正常。", AMBER_SOFT, AMBER)
    doc.add_heading("本版可演示内容", level=2)
    add_body(doc, "模拟病历展示、AI/安全回退选项、动态追问或直接建单、置信度处理、护理任务接单与处置记录。")
    page_break(doc)

    # Page 3: patient home.
    doc.add_heading("2  患者端界面", level=1)
    add_body(doc, "患者端是需求表达入口。演示人员可以点击候选项，也可以使用数字键 1～4 模拟 SSVEP 分类结果。")
    add_screenshot(doc, "v3-patient-home.png", "图 1  患者端一级安全菜单", "脑护通第三版患者端一级需求分类界面")
    for text in (
        "顶部：切换患者端/护理端，并显示当前病区和时间。",
        "中间：显示候选目标、当前步骤和已确认路径。",
        "右侧：切换模拟病历、调整置信度并查看事件记录。",
        "候选项下方：明确标注“继续确认”或“确认建单”。",
    ):
        add_list_item(doc, text, numbering["bullet"])
    page_break(doc)

    # Page 4: chart.
    doc.add_heading("3  查看模拟病历", level=1)
    add_body(doc, "在患者端右侧选择床位后，点击“查看完整模拟病历”。护理端任务详情中的“查看病历”也会打开同一份场景资料。")
    add_screenshot(doc, "v3-chart-drawer.png", "图 2  模拟病历抽屉", "床位A01模拟病历抽屉，包含诊断、表达能力、风险和护理注意事项")
    doc.add_heading("病历字段如何参与判断", level=2)
    for text in (
        "B02：吞咽风险和饮水状态用于拦截直接饮水动作。",
        "C03：术后体位限制用于拦截具体卧位动作。",
        "A01：没有特殊风险记录时，按普通受控引导流程继续。",
    ):
        add_list_item(doc, text, numbering["bullet"])
    add_callout(doc, "边界", "页面内所有诊断、年龄、风险和护理要求均为模拟信息。", RED_SOFT, RED)
    page_break(doc)

    # Page 5: AI path.
    doc.add_heading("4  AI动态路径与安全回退", level=1)
    add_body(doc, "一级菜单保持固定。进入下一层后，服务端结合已选需求和模拟病历生成受控候选项，并冻结本轮选项凭证。")
    add_screenshot(doc, "v3-ai-swallowing.png", "图 3  B02吞咽风险下的受控路径", "吞咽风险场景下显示安全回退路径、病历依据和动态候选项")
    for text in (
        "DeepSeek 实时分析：真实API成功返回，AI只调整审核候选项的顺序和短标签。",
        "AI路径模拟：未调用真实API，用于离线演示动态路径。",
        "安全回退路径：AI超时、报错或返回不合规内容时使用审核白名单。",
        "需评估：病历安全规则已触发，任务必须交由护理人员确认。",
    ):
        add_list_item(doc, text, numbering["bullet"], bold_prefix=text.split("：")[0] + "：")
    add_callout(doc, "安全原则", "AI不能新增医疗意图、修改风险等级、给出诊断或直接执行设备动作。", AMBER_SOFT, AMBER)
    page_break(doc)

    # Page 6: request confirmation.
    doc.add_heading("5  确认并发送患者需求", level=1)
    add_body(doc, "当意图已经明确时，系统进入最终确认。B02选择饮水口腔后不会继续询问饮水动作，而是直接形成护理评估需求。")
    add_screenshot(doc, "v3-request-review.png", "图 4  最终需求确认", "吞咽风险需求最终确认页，显示安全规则、路径判断和整体置信度")
    for text in (
        "检查需求文字、优先级和安全提示。",
        "检查每层选择和对应置信度。",
        "点击“确认并发送需求”。",
        "发送成功后点击“查看护理端”。",
    ):
        add_list_item(doc, text, numbering["number"])
    doc.add_heading("置信度规则", level=2)
    add_body(doc, "低于0.70拒绝；0.70～0.85要求再次确认；不低于0.85接受。整体置信度取已确认层级中的最低值。")
    page_break(doc)

    # Page 7: nursing side.
    doc.add_heading("6  护理端查看任务", level=1)
    add_body(doc, "护理端左侧是任务队列，右侧显示所选任务的完整确认链、病历摘要、风险依据和处置区。")
    add_screenshot(doc, "v3-nursing-task.png", "图 5  护理任务队列与患者需求详情", "护理端任务队列以及B02吞咽风险任务详情")
    for text in (
        "优先级：高、中、普通，用于演示排序和视觉提示。",
        "来源：区分脑控确认、受控引导或AI引导。",
        "意图确认链：显示每层选择与置信度。",
        "风险提示与依据：显示模拟病历字段、安全规则和AI/回退说明。",
    ):
        add_list_item(doc, text, numbering["bullet"])
    page_break(doc)

    # Page 8: nursing actions.
    doc.add_heading("7  护理任务处置闭环", level=1)
    add_body(doc, "待接单任务先点击“接单并前往床位”。接单后可以选择完成、进一步评估、暂时无法完成或转交。")
    add_screenshot(doc, "v3-nursing-actions.png", "图 6  护理端处置操作", "护理端接单后的进一步评估、无法完成和标记完成按钮")
    for text in (
        "标记已完成：护理人员已完成床旁评估与处置。",
        "需进一步评估：任务进入评估状态，评估结束后仍可完成。",
        "暂时无法完成：记录阻塞状态，等待转交或后续处理。",
        "转交任务：记录转交申请，不自动改变临床处置结论。",
    ):
        add_list_item(doc, text, numbering["bullet"], bold_prefix=text.split("：")[0] + "：")
    add_callout(doc, "审计", "所有处置操作都会写入页面事件记录。当前Demo数据保存在服务进程内存中，重启后恢复默认数据。")
    page_break(doc)

    # Page 9: demo scripts.
    doc.add_heading("8  推荐演示脚本", level=1)
    doc.add_heading("场景一：吞咽风险", level=2)
    for text in (
        "切换到 B02 · 吞咽风险。",
        "选择“基本照护”。",
        "观察“饮水口腔需评估”和“确认建单”标签。",
        "发送任务，在护理端查看病历依据并完成接单处置。",
    ):
        add_list_item(doc, text, numbering["number"])
    doc.add_heading("场景二：术后体位限制", level=2)
    add_body(doc, "切换到 C03，选择“基本照护”后观察体位需求如何直接转护理评估，而不继续生成左侧卧、右侧卧等具体动作。")
    doc.add_heading("场景三：疼痛继续追问", level=2)
    add_body(doc, "切换到 A01，选择“基本照护 → 疼痛不适”，系统继续询问疼痛部位和性质；选择腹部持续重痛后生成高优先级任务。")
    doc.add_heading("场景四：置信度二次确认", level=2)
    add_body(doc, "将右侧置信度滑块调整到0.80，再选择候选项，页面会要求再次确认；调整到0.69时输入会被拒绝。")
    add_callout(doc, "建议", "每次正式演示前点击护理端“恢复演示数据”，并先验证本地网址。", TEAL_SOFT, TEAL)
    page_break(doc)

    # Page 10: troubleshooting and boundaries.
    doc.add_heading("9  常见问题与应用边界", level=1)
    doc.add_heading("页面打不开", level=2)
    for text in (
        "确认 BAT 黑色窗口仍然开启。",
        "优先测试 http://127.0.0.1:8000/。",
        "如果提示端口不可用，关闭旧的演示窗口后重新运行 BAT。",
        "公网地址失效时重新启动 cpolar，并以黑色窗口最新显示的网址为准。",
    ):
        add_list_item(doc, text, numbering["bullet"])
    doc.add_heading("AI没有生成或显示安全回退", level=2)
    add_body(doc, "检查 frontend/.env.local 中的AI模式、模型和API密钥。AI失败不会阻断演示，系统会自动使用审核白名单。")
    doc.add_heading("当前系统可以做什么", level=2)
    add_body(doc, "模拟患者脑控选择，结合模拟病历生成受控引导，执行置信度和风险判断，形成护理任务并记录护理处置。")
    doc.add_heading("当前系统不能做什么", level=2)
    for text in (
        "不能进行医学诊断、治疗决策或药物建议。",
        "不能替代护士床旁评估、医院呼叫系统或急救流程。",
        "尚未接入真实OpenBCI/SSVEP信号。",
        "尚未控制小米智能设备，设备接口保持关闭。",
        "尚未接入真实HIS、电子病历或生产数据库。",
    ):
        add_list_item(doc, text, numbering["bullet"])
    add_callout(doc, "最终责任边界", "系统负责表达、确认和传递；护理人员负责床旁评估、判断和干预。", RED_SOFT, RED)

    doc.core_properties.title = "脑护通Demo V3用户使用手册"
    doc.core_properties.subject = "AI临床情境辅助受控照护交互演示操作说明"
    doc.core_properties.author = "脑护通项目组"
    doc.core_properties.keywords = "脑护通, Demo V3, 用户手册, 脑控表达, 护理任务"
    doc.save(OUTPUT)
    structural_audit(OUTPUT)
    print(OUTPUT)


def structural_audit(path):
    doc = Document(path)
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    required = [
        "脑护通 Demo V3",
        "快速启动",
        "模拟病历",
        "AI动态路径",
        "护理任务处置闭环",
        "常见问题与应用边界",
    ]
    for item in required:
        if item not in text:
            raise AssertionError(f"missing required text: {item}")
    section = doc.sections[0]
    if abs(section.page_width.inches - 8.5) > 0.01 or abs(section.page_height.inches - 11) > 0.01:
        raise AssertionError("page size is not US Letter")
    for margin in (section.top_margin, section.right_margin, section.bottom_margin, section.left_margin):
        if abs(margin.inches - 1) > 0.01:
            raise AssertionError("page margins do not match compact_reference_guide")
    with ZipFile(path) as archive:
        media = [name for name in archive.namelist() if name.startswith("word/media/")]
        if len(media) < 6:
            raise AssertionError("expected at least six screenshots")


if __name__ == "__main__":
    build_manual()
