from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
ASSETS = Path(__file__).resolve().parent / "assets"
OUTPUT = ROOT / "docs" / "脑护通Demo客户介绍手册.docx"

TEAL = "157B78"
TEAL_DARK = "105E5B"
TEAL_LIGHT = "EAF5F4"
INK = "172221"
MUTED = "5F6B69"
LINE = "D5DEDC"
SOFT = "F4F8F7"
RED = "B52235"
RED_LIGHT = "FCEDEF"
AMBER = "A46808"
AMBER_LIGHT = "FFF6DF"
WHITE = "FFFFFF"


def set_run_font(run, size=None, bold=None, color=None, name="Microsoft YaHei"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=120, start=150, bottom=120, end=150):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color=LINE, size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "nil")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, 8, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.addnext(fld)
    tail = paragraph.add_run(" 页")
    set_run_font(tail, 8, color=MUTED)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(15)
    section.bottom_margin = Mm(14)
    section.left_margin = Mm(17)
    section.right_margin = Mm(17)
    section.header_distance = Mm(7)
    section.footer_distance = Mm(7)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color in (
        ("Heading 1", 23, INK),
        ("Heading 2", 15, TEAL_DARK),
        ("Heading 3", 11, INK),
    ):
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True

    header = section.header
    table = header.add_table(rows=1, cols=2, width=Mm(176))
    remove_table_borders(table)
    table.columns[0].width = Mm(105)
    table.columns[1].width = Mm(71)
    left = table.cell(0, 0).paragraphs[0]
    left.paragraph_format.space_after = Pt(0)
    set_run_font(left.add_run("脑护通"), 9, True, TEAL_DARK)
    set_run_font(left.add_run("  |  DEMO CLIENT GUIDE"), 8, False, MUTED)
    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right.paragraph_format.space_after = Pt(0)
    set_run_font(right.add_run("受控照护交互演示"), 8, False, MUTED)

    footer = section.footer
    ftable = footer.add_table(rows=1, cols=2, width=Mm(176))
    remove_table_borders(ftable)
    left_footer = ftable.cell(0, 0).paragraphs[0]
    left_footer.paragraph_format.space_after = Pt(0)
    set_run_font(left_footer.add_run("客户演示版  |  2026.07"), 8, False, MUTED)
    add_page_number(ftable.cell(0, 1).paragraphs[0])


def add_kicker(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    set_run_font(p.add_run(text.upper()), 8, True, TEAL_DARK)
    return p


def add_title(doc, title, subtitle=None):
    add_kicker(doc, "BRAIN CARE DEMO")
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    set_run_font(p.add_run(title), 23, True, INK)
    if subtitle:
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(12)
        set_run_font(sp.add_run(subtitle), 10.5, False, MUTED)


def add_rule(doc, color=TEAL):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(12)
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    p_pr.append(borders)


def add_callout(doc, title, body, fill=TEAL_LIGHT, accent=TEAL_DARK):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Mm(176)
    cell = table.cell(0, 0)
    cell.width = Mm(176)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 180, 220, 180, 220)
    set_cell_border(cell, color=accent, size="10")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    set_run_font(p.add_run(title), 10.5, True, accent)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.2
    set_run_font(p2.add_run(body), 9.5, False, INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_feature_cards(doc, cards, columns=2):
    rows = (len(cards) + columns - 1) // columns
    table = doc.add_table(rows=rows, cols=columns)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    remove_table_borders(table)
    widths = [Mm(86)] * columns
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = widths[index]
            set_cell_margins(cell, 130, 130, 130, 130)
    for index, (number, title, body) in enumerate(cards):
        cell = table.cell(index // columns, index % columns)
        set_cell_shading(cell, SOFT)
        set_cell_border(cell, LINE, "6")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(3)
        set_run_font(p.add_run(number), 9, True, TEAL_DARK)
        set_run_font(p.add_run(f"  {title}"), 11, True, INK)
        b = cell.add_paragraph()
        b.paragraph_format.space_after = Pt(0)
        set_run_font(b.add_run(body), 9, False, MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_steps_table(doc, steps):
    table = doc.add_table(rows=1, cols=len(steps))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    remove_table_borders(table)
    width = Mm(176 / len(steps))
    for index, (number, title, body) in enumerate(steps):
        cell = table.cell(0, index)
        cell.width = width
        set_cell_shading(cell, TEAL_LIGHT if index % 2 == 0 else SOFT)
        set_cell_border(cell, LINE, "6")
        set_cell_margins(cell, 160, 130, 160, 130)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        set_run_font(p.add_run(number), 12, True, TEAL_DARK)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(3)
        set_run_font(p2.add_run(title), 10, True, INK)
        p3 = cell.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.paragraph_format.space_after = Pt(0)
        set_run_font(p3.add_run(body), 8.5, False, MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_bullets(doc, items, color=INK):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Mm(5)
        p.paragraph_format.first_line_indent = Mm(-2)
        p.paragraph_format.space_after = Pt(4)
        set_run_font(p.add_run(item), 9.5, False, color)


def add_picture(doc, path, width_mm=176, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    run.add_picture(str(path), width=Mm(width_mm))
    doc_pr = run._r.xpath(".//wp:docPr")
    if doc_pr:
        doc_pr[0].set("descr", caption or path.stem)
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(7)
        set_run_font(cp.add_run(caption), 8.5, False, MUTED)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def build_document():
    doc = Document()
    configure_document(doc)

    # Cover
    doc.add_paragraph().paragraph_format.space_after = Pt(30)
    tag = doc.add_paragraph()
    tag.paragraph_format.space_after = Pt(12)
    set_run_font(tag.add_run("脑  护  通"), 12, True, TEAL_DARK)
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(8)
    set_run_font(title.add_run("脑护通 Demo"), 34, True, INK)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    set_run_font(subtitle.add_run("客户介绍手册"), 23, True, TEAL_DARK)
    lead = doc.add_paragraph()
    lead.paragraph_format.space_after = Pt(22)
    lead.paragraph_format.line_spacing = 1.35
    set_run_font(lead.add_run("受控照护交互与脑控需求闭环演示"), 14, False, MUTED)
    add_rule(doc)
    promise = doc.add_paragraph()
    promise.paragraph_format.space_after = Pt(22)
    promise.paragraph_format.line_spacing = 1.4
    set_run_font(
        promise.add_run("让受限表达患者通过分层选择，把需求清晰、可靠、可追溯地送达护理端。"),
        17,
        True,
        INK,
    )
    add_callout(
        doc,
        "演示范围",
        "当前版本用于展示交互流程和安全机制。页面中的床位、护理需求和事件均为虚构演示数据，不用于医学诊断或真实临床处置。",
        fill=SOFT,
        accent=TEAL,
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(24)
    meta = doc.add_table(rows=2, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    meta.autofit = False
    remove_table_borders(meta)
    values = (("版本", "Demo 0.1"), ("发布日期", "2026 年 7 月"))
    for row_index, (label, value) in enumerate(values):
        meta.cell(row_index, 0).width = Mm(35)
        meta.cell(row_index, 1).width = Mm(90)
        p1 = meta.cell(row_index, 0).paragraphs[0]
        p2 = meta.cell(row_index, 1).paragraphs[0]
        set_run_font(p1.add_run(label), 9, True, MUTED)
        set_run_font(p2.add_run(value), 9, False, INK)

    # Page 2
    add_page_break(doc)
    add_title(doc, "产品定位", "面向受限表达场景的需求表达与护理任务闭环 Demo")
    add_callout(
        doc,
        "一句话理解",
        "患者通过分层候选项表达需求，系统对每一步置信度进行安全判断，再将完整意图链路交付护理端处理。",
    )
    doc.add_heading("当前 Demo 解决什么问题", level=2)
    add_feature_cards(
        doc,
        [
            ("01", "表达更具体", "通过三层选择，把“我不舒服”逐步澄清为可执行护理需求。"),
            ("02", "结果更稳妥", "低置信度拒绝、中间区间二次确认，高置信度正常接受。"),
            ("03", "任务可追踪", "护理端查看床位、需求、来源、优先级、状态和确认链路。"),
            ("04", "流程可扩展", "后续可接入真实 SSVEP 分类器、数据库以及 AI 动态引导选项。"),
        ],
    )
    doc.add_heading("端到端演示流程", level=2)
    add_steps_table(
        doc,
        [
            ("1", "选择类型", "疼痛、呼吸、饮水、体位"),
            ("2", "逐层澄清", "确认部位、时效或具体诉求"),
            ("3", "安全确认", "逐层判断置信度"),
            ("4", "任务闭环", "创建、接单、完成并留痕"),
        ],
    )
    doc.add_heading("适合展示的场景", level=2)
    add_bullets(
        doc,
        [
            "脑机接口、无障碍交互、智慧病房和护理信息化方案演示。",
            "客户需求沟通、概念验证、比赛答辩及内部产品评审。",
            "在接入真实硬件前，先验证患者表达和护理处理流程。",
        ],
    )

    # Page 3
    add_page_break(doc)
    add_title(doc, "患者端：分层表达", "大尺寸候选项适合注视目标，也支持点击和数字键完成当前 Demo 操作")
    add_picture(doc, ASSETS / "patient-view.png", 176, "患者端首页：一级需求选择与模拟脑控信号面板")
    doc.add_heading("演示操作", level=2)
    add_steps_table(
        doc,
        [
            ("1", "选择需求", "选择疼痛不适等一级分类"),
            ("2", "确认细节", "继续选择部位和程度"),
            ("3", "查看链路", "核对三层结果及置信度"),
            ("4", "发送需求", "确认后进入护理任务队列"),
        ],
    )
    add_callout(
        doc,
        "真实设备接入方式",
        "当前使用点击或数字键模拟 SSVEP 分类结果。未来真实分类器只需把候选项编号与置信度提交到相同接口，页面和护理流程无需重做。",
        fill=SOFT,
        accent=TEAL,
    )

    # Page 4
    add_page_break(doc)
    add_title(doc, "置信度安全机制", "置信度用于判断本次脑控选择是否足够可靠，不代表医学诊断结论")
    table = doc.add_table(rows=4, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    headers = ("置信度区间", "系统动作", "患者体验", "目的")
    for index, text in enumerate(headers):
        cell = table.cell(0, index)
        set_cell_shading(cell, TEAL_DARK)
        set_cell_margins(cell, 140, 120, 140, 120)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.add_run(text), 9, True, WHITE)
    rows = (
        ("低于 0.70", "拒绝输入", "保持当前页面，可重新选择", "避免明显不可靠输入"),
        ("0.70 至 0.85", "再次确认", "显示待确认候选项", "降低误触发风险"),
        ("不低于 0.85", "正常接受", "进入下一层或最终确认", "兼顾效率与可靠性"),
    )
    fills = (RED_LIGHT, AMBER_LIGHT, TEAL_LIGHT)
    for row_index, values in enumerate(rows, start=1):
        for col_index, value in enumerate(values):
            cell = table.cell(row_index, col_index)
            set_cell_shading(cell, fills[row_index - 1] if col_index == 0 else WHITE)
            set_cell_border(cell, LINE, "6")
            set_cell_margins(cell, 150, 120, 150, 120)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_index < 2 else WD_ALIGN_PARAGRAPH.LEFT
            color = (RED, AMBER, TEAL_DARK)[row_index - 1] if col_index == 0 else INK
            set_run_font(p.add_run(value), 9, col_index == 0, color)
    doc.add_paragraph()
    add_callout(
        doc,
        "整体可靠度",
        "最终页面展示三层置信度，例如 0.91 / 0.88 / 0.93。整体可靠度取其中最低值 0.88，避免某一层较弱结果被平均值掩盖。",
    )
    doc.add_heading("安全边界", level=2)
    add_bullets(
        doc,
        [
            "置信度只衡量分类器对候选项的确定程度，不判断患者病情。",
            "护理任务的最终执行由护理人员确认，系统不自动进行医学处置。",
            "未来接入 AI 时，AI只生成受控引导选项，不能修改脑控阈值或替代人工确认。",
        ],
    )

    # Page 5
    add_page_break(doc)
    add_title(doc, "护理端：任务闭环", "护理人员在同一页面查看需求、置信度链路、优先级和处理状态")
    add_picture(doc, ASSETS / "care-detail-view.png", 176, "护理端：任务队列与患者需求详情")
    doc.add_heading("护理端可见信息", level=2)
    add_feature_cards(
        doc,
        [
            ("A", "任务摘要", "床位、需求内容、来源和优先级。"),
            ("B", "意图链路", "每层选择结果及对应置信度。"),
            ("C", "处理状态", "待接单、已接单、已完成。"),
            ("D", "事件记录", "任务创建和状态变化均保留时间记录。"),
        ],
    )
    add_callout(
        doc,
        "状态约束",
        "任务按照“待接单 → 已接单 → 已完成”顺序推进。非法跳转会被后端拒绝，避免界面状态与实际流程不一致。",
        fill=SOFT,
        accent=TEAL,
    )

    # Page 6
    add_page_break(doc)
    add_title(doc, "当前边界与升级路线", "本版本先验证流程，再按项目阶段逐步接入真实能力")
    doc.add_heading("当前版本说明", level=2)
    boundary = doc.add_table(rows=4, cols=2)
    boundary.alignment = WD_TABLE_ALIGNMENT.CENTER
    boundary.autofit = False
    limits = (
        ("输入信号", "点击或数字键模拟脑控选择；真实 SSVEP 分类器接口已预留。"),
        ("数据存储", "任务和事件保存在服务进程内存中；重启服务后恢复默认演示数据。"),
        ("身份权限", "当前公开 Demo 未接入登录和角色权限，仅可使用虚构演示数据。"),
        ("临床用途", "当前版本不用于诊断、治疗建议或真实临床处置。"),
    )
    for row_index, (label, body) in enumerate(limits):
        left = boundary.cell(row_index, 0)
        right = boundary.cell(row_index, 1)
        left.width = Mm(36)
        right.width = Mm(140)
        set_cell_shading(left, SOFT)
        set_cell_shading(right, WHITE)
        set_cell_border(left, LINE, "6")
        set_cell_border(right, LINE, "6")
        set_cell_margins(left, 145, 140, 145, 140)
        set_cell_margins(right, 145, 160, 145, 160)
        set_run_font(left.paragraphs[0].add_run(label), 9, True, TEAL_DARK)
        set_run_font(right.paragraphs[0].add_run(body), 9, False, INK)
    doc.add_heading("建议升级顺序", level=2)
    add_steps_table(
        doc,
        [
            ("1", "动态引导", "AI根据已选分类生成受控二级选项"),
            ("2", "真实脑控", "接入 SSVEP 分类结果和设备状态"),
            ("3", "持久化", "数据库、登录、权限和审计日志"),
            ("4", "系统集成", "对接护理工作站或院内任务系统"),
        ],
    )
    doc.add_heading("AI升级原则", level=2)
    add_bullets(
        doc,
        [
            "固定一级安全分类，AI只生成简短、互斥、可回退的二级引导选项。",
            "紧急求助等关键入口保持固定，不依赖 AI 在线状态。",
            "AI调用失败时自动使用预设选项，不阻断患者表达。",
            "护理人员保留最终确认权，AI不自动诊断、接单或执行任务。",
        ],
    )
    add_callout(
        doc,
        "客户演示建议",
        "建议按“患者端完成三层选择 → 查看最终置信度 → 发送需求 → 切换护理端接单并完成”的顺序演示，完整流程约 3 分钟。",
        fill=TEAL_LIGHT,
        accent=TEAL_DARK,
    )

    doc.core_properties.title = "脑护通 Demo 客户介绍手册"
    doc.core_properties.subject = "受控照护交互与脑控需求闭环演示"
    doc.core_properties.author = "脑护通项目组"
    doc.core_properties.keywords = "脑护通, 脑控交互, 护理任务, Demo"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
